from datetime import datetime
import os
import re
import tempfile
import asyncio
import mimetypes
import time
import requests
from voice_config import VOICES
from docx import Document as DocxDocument
from dotenv import load_dotenv
from typing import List, Union, Dict, Any
from pydantic import BaseModel
from langchain_community.document_loaders import PyPDFLoader, CSVLoader, UnstructuredExcelLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain.schema import Document
import aiohttp
from crawl4ai import AsyncWebCrawler
from crawl4ai.async_configs import BrowserConfig, CrawlerRunConfig
from moviepy.editor import VideoFileClip
from urllib.parse import urlparse
from models import ModelManager
import logging
import json

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('document_processor.log'),
        logging.StreamHandler()
    ]
)

load_dotenv()

class ProcessingResult(BaseModel):
    success: bool
    content: str
    metadata: Dict[str, Any] = {}
    error_message: str = ""

class DocumentProcessor:
    def __init__(self):
        self.headers = {"Authorization": f"Bearer {os.getenv('HF_API_KEY')}"}
        if not self.headers["Authorization"]:
            raise ValueError("HF_API_KEY not found in environment variables")
            
        self.browser_config = BrowserConfig(verbose=False)
        self.crawler_run_config = CrawlerRunConfig(
            word_count_threshold=50,
            remove_overlay_elements=True,
            process_iframes=True
        )
        self.MAX_FILE_SIZE = 25 * 1024 * 1024  # 25MB for audio/video
        self.MAX_AUDIO_LENGTH = 300  # 5 minutes in seconds
        
        # Create tmp directory for processing
        self.tmp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tmp")
        os.makedirs(self.tmp_dir, exist_ok=True)
        
        # Initialize supported mime types
        self.supported_mime_types = {
            'pdf': 'application/pdf',
            'csv': 'text/csv',
            'excel': 'application/vnd.ms-excel',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'doc': 'application/msword',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain',
            'image': 'image/',
            'video': 'video/',
            'audio': 'audio/'
        }

    async def _make_api_request(self, url: str, data: bytes, timeout: int = 300) -> Dict:
        async with aiohttp.ClientSession() as session:
            try:
                async with session.post(url, headers=self.headers, data=data, timeout=timeout) as response:
                    if response.status == 413:
                        raise ValueError("File is too large for API processing")
                    return await response.json()
            except asyncio.TimeoutError:
                raise TimeoutError("API request timed out")

    def _check_file_size(self, file_path: str) -> bool:
        if not os.path.isfile(file_path):
            return True
        
        mime_type = self._get_mime_type(file_path)
        if mime_type and (mime_type.startswith('audio/') or mime_type.startswith('video/')):
            size_limit = self.MAX_FILE_SIZE
        else:
            size_limit = self.MAX_FILE_SIZE * 2
            
        file_size = os.path.getsize(file_path)
        return file_size <= size_limit

    def _get_mime_type(self, file_path: str) -> str:
        mime_type = mimetypes.guess_type(file_path)[0]
        if mime_type is None:
            # Try to determine type from file extension
            ext = os.path.splitext(file_path)[1].lower()
            mime_map = {
                '.pdf': 'application/pdf',
                '.csv': 'text/csv',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel',
                '.doc': 'application/msword',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.txt': 'text/plain',
                '.mp3': 'audio/mpeg',
                '.wav': 'audio/wav',
                '.mp4': 'video/mp4',
                '.avi': 'video/x-msvideo'
            }
            mime_type = mime_map.get(ext)
        return mime_type

    async def process_file(self, file_path: str) -> ProcessingResult:
        try:
            if urlparse(file_path).scheme in ('http', 'https'):
                result = await self._process_url(file_path)
                logging.info(f"Successfully processed URL: {file_path}")
                return ProcessingResult(
                    success=True,
                    content=result,
                    metadata={'type': 'url', 'path': file_path}
                )
            
            if not os.path.exists(file_path):
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"File not found: {file_path}"
                )

            if not self._check_file_size(file_path):
                msg = f"File {file_path} exceeds size limit"
                logging.warning(msg)
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=msg
                )

            mime_type = self._get_mime_type(file_path)
            if not mime_type:
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"Unknown file type: {file_path}"
                )

            logging.info(f"Processing {file_path} ({mime_type})")

            result = None
            metadata = {'type': mime_type, 'path': file_path}

            if mime_type == "application/pdf":
                result = self._process_pdf(file_path)
            elif mime_type in ["text/csv", "application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                result = self._process_csv_excel(file_path)
            elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                result = self._process_doc(file_path)
            elif mime_type and mime_type.startswith('image/'):
                result = await self._process_image(file_path)
            elif mime_type and mime_type.startswith('video/'):
                result = await self._process_video(file_path)
            elif mime_type and mime_type.startswith('audio/'):
                result = await self._process_audio(file_path)
            elif mime_type == "text/plain":
                with open(file_path, 'r', encoding='utf-8') as f:
                    result = f.read()
            else:
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"Unsupported file type: {mime_type}"
                )

            if result:
                logging.info(f"Successfully processed: {file_path}")
                return ProcessingResult(
                    success=True,
                    content=result,
                    metadata=metadata
                )
            else:
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"Failed to process {os.path.basename(file_path)}"
                )

        except Exception as e:
            logging.error(f"Error processing {file_path}: {str(e)}", exc_info=True)
            return ProcessingResult(
                success=False,
                content="",
                error_message=f"Error processing {os.path.basename(file_path)}: {str(e)}"
            )

    def _process_pdf(self, path: str) -> str:
        try:
            loader = PyPDFLoader(path)
            pages = loader.load()
            return "\n".join([p.page_content for p in pages])
        except Exception as e:
            raise RuntimeError(f"Error processing PDF file: {str(e)}")

    def _process_csv_excel(self, path: str) -> str:
        try:
            if path.endswith('.csv'):
                loader = CSVLoader(path)
            else:
                loader = UnstructuredExcelLoader(path)
            docs = loader.load()
            return "\n".join([d.page_content for d in docs])
        except Exception as e:
            raise RuntimeError(f"Error processing CSV/Excel file: {str(e)}")

    def _process_doc(self, path: str) -> str:
        try:
            doc = DocxDocument(path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    full_text.append(para.text)
            # Also process tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            return "\n".join(full_text)
        except Exception as e:
            raise RuntimeError(f"Error processing document file: {str(e)}")

    async def _process_image(self, path: str) -> str:
        API_URL = "https://api-inference.huggingface.co/models/Salesforce/blip-image-captioning-large"
        try:
            with open(path, "rb") as f:
                data = f.read()
            result = await self._make_api_request(API_URL, data)
            return result[0]['generated_text'] if isinstance(result, list) else result.get('error', 'Image processing failed')
        except Exception as e:
            raise RuntimeError(f"Error processing image file: {str(e)}")

    async def _process_video(self, path: str) -> str:
        try:
            video = VideoFileClip(path)
            if video.duration > self.MAX_AUDIO_LENGTH:
                video.close()
                return f"Error: Video duration exceeds {self.MAX_AUDIO_LENGTH} seconds limit"
            
            with tempfile.NamedTemporaryFile(suffix=".wav", dir=self.tmp_dir, delete=False) as tmpfile:
                video.audio.write_audiofile(tmpfile.name, logger=None)
                result = await self._process_audio(tmpfile.name)
                video.close()
                
                try:
                    os.unlink(tmpfile.name)
                except Exception as e:
                    logging.warning(f"Failed to delete temporary file {tmpfile.name}: {str(e)}")
                    
                return f"Video transcription:\n{result}"
        except Exception as e:
            raise RuntimeError(f"Error processing video file: {str(e)}")

    async def _process_audio(self, path: str) -> str:
        API_URL = "https://api-inference.huggingface.co/models/openai/whisper-large-v3-turbo"
        try:
            if os.path.getsize(path) > self.MAX_FILE_SIZE:
                return "Error: Audio file is too large for processing"
                
            with open(path, "rb") as f:
                data = f.read()
            result = await self._make_api_request(API_URL, data)
            transcription = result.get('text', '')
            if not transcription:
                raise ValueError("No transcription received from API")
            return f"Audio transcription:\n{transcription}"
        except Exception as e:
            raise RuntimeError(f"Error processing audio file: {str(e)}")

    async def _process_url(self, path: str) -> str:
        try:
            async with AsyncWebCrawler(config=self.browser_config) as crawler:
                result = await crawler.arun(
                    url=path,
                    config=self.crawler_run_config
                )
                if result.success:
                    return result.markdown
                else:
                    raise RuntimeError("Failed to crawl URL")
        except Exception as e:
            raise RuntimeError(f"Error processing URL: {str(e)}")

class RAGSystem:
    def __init__(self, model_name: str = "llama-3.3-70b-versatile"):
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        self.vector_store = None
        self.model_manager = ModelManager()
        self.llm = self.model_manager.get_model(model_name)
        self.document_metadata = []

    async def ingest_documents(self, documents: List[ProcessingResult]):
        if not documents:
            raise ValueError("No documents provided for ingestion")
        
        processed_docs = []
        for doc in documents:
            if doc.success and doc.content:
                doc_chunks = self.text_splitter.split_text(doc.content)
                processed_docs.extend([
                    Document(
                        page_content=chunk,
                        metadata=doc.metadata
                    ) for chunk in doc_chunks
                ])
                self.document_metadata.append(doc.metadata)
        
        if not processed_docs:
            raise ValueError("No valid documents to process")
        
        if self.vector_store is None:
            self.vector_store = Chroma.from_documents(processed_docs, self.embeddings)
        else:
            self.vector_store.add_documents(processed_docs)

    def get_ingested_documents_info(self) -> str:
        if not self.document_metadata:
            return "No documents have been ingested yet."
        
        info = ["Ingested documents:"]
        for meta in self.document_metadata:
            doc_type = meta.get('type', 'unknown')
            path = meta.get('path', 'unknown')
            info.append(f"- {os.path.basename(path)} (Type: {doc_type})")
        return "\n".join(info)

    async def query(self, question: str, k: int = 5) -> str:
        if self.vector_store is None:
            return "No documents loaded yet. Please ingest documents first."
        
        docs = self.vector_store.similarity_search(question, k=k)
        context = "\n\n".join([
            f"[Source: {os.path.basename(d.metadata.get('path', 'unknown'))}]\n{d.page_content}"
            for d in docs
        ])
        
        prompt = ChatPromptTemplate.from_template(
            """
            ### Instructions
            You are an intelligent RAG-based chatbot designed to assist users by providing detailed and relevant answers based on the provided context. 
            You will answer questions by retrieving the most relevant documents and generating responses based solely on the retrieved information.
            When referencing information, mention the source document in your response.

            ### Context
            {context}

            ### User's Question
            {question}

            ### Response
            Please provide a comprehensive and well-structured response to the user's question based on the context. 
            Ensure your answer is accurate, informative, and properly cites the sources used.
            If the question cannot be fully answered using the provided context, acknowledge this limitation.

            ---
            """
        )
        
        try:
            chain = prompt | self.llm
            response = await chain.ainvoke({"context": context, "question": question})
            return response.content
        except Exception as e:
            logging.error(f"Error during query processing: {str(e)}", exc_info=True)
            return f"Error processing query: {str(e)}"

async def process_files_with_progress(processor: DocumentProcessor, file_paths: List[str]) -> List[ProcessingResult]:
    total_files = len(file_paths)
    processed = 0
    results = []
    
    print(f"\nProcessing {total_files} files...")
    
    for path in file_paths:
        processed += 1
        print(f"Processing ({processed}/{total_files}): {os.path.basename(path)}")
        result = await processor.process_file(path)
        results.append(result)
        
        if result.success:
            print(f"✓ Success: {os.path.basename(path)}")
        else:
            print(f"✗ Failed: {os.path.basename(path)} - {result.error_message}")
    
    return results

class PodcastProcessor:
    def __init__(self, user_id: str, secret_key: str):
        self.model_manager = ModelManager()
        self.llm = self.model_manager.get_model("deepseek-r1-distill-llama-70b")
        self.user_id = user_id
        self.secret_key = secret_key
        self.headers = {
            'X-USER-ID': user_id,
            'Authorization': f'Bearer {secret_key}',
            'Content-Type': 'application/json',
        }

    async def process_podcast(self, document_content: str) -> str:
        """Generate podcast transcript from document content"""
        prompt = ChatPromptTemplate.from_template(
            """
            ### Podcast Dialogue Generator
            Create a compelling podcast conversation between two hosts that transforms complex content into engaging dialogue. Follow these guidelines:

            ## Host Personas
            1. **Dynamic Duo:**
            - Host 1: "The Enthusiast" - Curious, energetic, asks probing questions
            - Host 2: "The Expert" - Knowledgeable, witty, delivers key insights
            - Natural back-and-forth flow without interruptions

            ## Content Requirements
            1. **Conversation Style:**
            - Pure dialogue format: "Host 1:..." "Host 2:..." 
            - No segment titles, music cues, or formatting
            - Seamless topic transitions through dialogue

            2. **Content Adaptation:**
            - Convert technical terms using relatable analogies
            - Use "Imagine if..." explanations for complex ideas
            - Highlight 2-3 key insights from document

            ## Dialogue Rules
            1. **Flow & Structure:**
            - Alternate every 1-3 sentences
            - Use conversational connectors: "Right...", "But consider...", "Here's why..."
            - Include 3 audience engagement phrases per 500 words: "Ever wondered...", "Picture this..."
            - Create engaging and dependend sentences and also add human like interactions like hmm, okay, right, aah, got it, etc.

            2. **Tone Balance:**
            - 30 percent humor/references, 50 percent insights, 20 percent banter
            - Professional foundation with witty spikes
            - Example: "So it's like TikTok for neurons? (laughs) But seriously..."

            ## Technical Specs
            - Length: you are technically free to make it as long as you want but try to keep it around 12 mins to 22 mins
            - Complexity: Grade 8-10 reading level
            - Format: Strictly "Host 1: [text]" lines without empty lines

            ## Required Content
            {document_content}

            ## Anti-Requirements
            - No markdown/formatting
            - No section headers or labels
            - No monologues >3 sentences
            - No passive voice

            ## Example Format:
            Host 1: Welcome to The Tech Tomorrow Podcast! Today we're diving into AI voice technology.
            Host 2: And what a topic this is. The progress from basic commands to full conversations is staggering.
            Host 1: Remember when asking phones to set timers felt revolutionary?
            Host 2: Now AI understands context and nuance. But does that raise ethical questions?
            Host 1: Exactly! Where should we draw the privacy line with these voice assistants?

            ## Example format2:
            Speaker 1: Alright so, instead of our regularly scheduled programming, there's something hairier that came across my feed last night that I just need to discuss.
            Speaker 2: Wait, where is this going?
            Speaker 1: I uh, I just thought we'd take a little... detour. You know, take the scenic route down the path of mystery. Specifically, into the thick, mossy woods where something like, oh I don't know… *Bigfoot* himself might be lurking.
            Speaker 2: Oh, for the love of, again, Briggs? Really? We did this last month. And the month before that. This is basically the podcast equivalent of your sad karaoke go-to.
            Speaker 1: What? No! I'm just... providing the people what they want! Listen, There's new evidence, and the public demands more attention on it.
            Speaker 2: The "public" is just you, Briggs. You're the one emailing us suggestions under fake names. We've all seen "Biggie O Footlore" in our inbox.
            """
        )

        
        chain = prompt | self.llm
        response = await chain.ainvoke({"document_content": document_content})
        transcript = response.content
        
        self._save_transcript(transcript)
        return transcript

    def _save_transcript(self, transcript: str):
        """Save generated transcript to transcripts folder"""
        os.makedirs("transcripts", exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"transcript_{timestamp}.txt"
        
        # Extract only the host dialogues
        host_transcript = []
        in_think_tag = False
        for line in transcript.splitlines():
            if "<think>" in line:
                in_think_tag = True
            elif "</think>" in line:
                in_think_tag = False
            elif not in_think_tag:
                host_transcript.append(line)
        
        host_transcript = "\n".join(host_transcript)
        
        with open(f"transcripts/{filename}", "w") as f:
            f.write(host_transcript)
        print(f"Transcript saved to transcripts/{filename}")

    def generate_audio(self, transcript: str, host1_voice: str, host2_voice: str) -> str:
        """Convert transcript to audio using specified voices"""
        # Get voice IDs from configuration
        voice1_id = VOICES.get(host1_voice, {}).get("id")
        voice2_id = VOICES.get(host2_voice, {}).get("id")
        
        if not voice1_id or not voice2_id:
            raise ValueError("Invalid voice selection")

        payload = {
            'model': 'PlayDialog',
            'text': transcript,
            'voice': voice1_id,
            'voice2': voice2_id,
            'turnPrefix': 'Host 1:',
            'turnPrefix2': 'Host 2:',
            'outputFormat': 'mp3',
        }

        response = requests.post(
            'https://api.play.ai/api/v1/tts/',
            headers=self.headers,
            json=payload
        )

        if response.status_code != 201:
            raise Exception(f"TTS API error: {response.text}")

        job_id = response.json()['id']
        return self._poll_audio_job(job_id)

    def _poll_audio_job(self, job_id: str) -> str:
        """Poll audio generation job until completion"""
        url = f'https://api.play.ai/api/v1/tts/{job_id}'
        while True:
            response = requests.get(url, headers=self.headers)
            if not response.ok:
                raise Exception(f"Status check failed: {response.text}")
                
            status = response.json().get('output', {}).get('status')
            if status == 'COMPLETED':
                return response.json()['output']['url']
            if status == 'FAILED':
                raise Exception("Audio generation failed")
                
            time.sleep(5)


async def main():
    try:
        processor = DocumentProcessor()
        model_manager = ModelManager()
        # Initialize PodcastProcessor with your credentials
        podcast_processor = PodcastProcessor(
            user_id="ssFAgjpDO5NVuWyC9AOLghsnTlS2",
            secret_key="ak-f16737b738ea4dd78ea3aefa81fb9c25"
        )
        
        print("Available Models:")
        for model in model_manager.list_models():
            print(f"- {model['name']} ({model['provider']}) - Cost: ${model['cost_per_million_tokens']} per million tokens")
        
        while True:
            model_name = input("\nEnter model name (or press Enter for default llama-3.3-70b-versatile): ").strip()
            if not model_name:
                model_name = "llama-3.3-70b-versatile"
            try:
                rag = RAGSystem(model_name)
                break
            except ValueError as e:
                print(f"Error: {e}")
        
        # Process files
        file_paths = [
            # "https://docs.crawl4ai.com/core/installation/",
            # "https://docs.crawl4ai.com/core/fit-markdown/",
            "https://genius.com/Lord-huron-mine-forever-lyrics",
            # "./docs/LLM_Example.doc",
            "./docs/video.mp4",
            "./docs/audio.mp3"
        ]

        # Allow user to add more files
        while True:
            additional_file = input("\nEnter additional file path (or press Enter to continue): ").strip()
            if not additional_file:
                break
            file_paths.append(additional_file)

        # Process files with progress tracking
        results = await process_files_with_progress(processor, file_paths)
        
        # Filter out failed results and ingest successful ones
        successful_results = [r for r in results if r.success]
        failed_results = [r for r in results if not r.success]
        
        if failed_results:
            print("\nFailed to process the following files:")
            for result in failed_results:
                print(f"- {result.metadata.get('path', 'unknown')}: {result.error_message}")
        
        if successful_results:
            print("\nIngesting successfully processed documents...")
            await rag.ingest_documents(successful_results)
            print("\n" + rag.get_ingested_documents_info())
        else:
            print("\nNo documents were successfully processed for ingestion.")
            return

        # Chat interface
        print("\nRAG System Ready. Available commands:")
        print("- 'exit': Quit the program")
        print("- 'switch model': Change the language model")
        print("- 'info': Show ingested documents information")
        print("- 'create podcast': Generate a podcast from ingested content")
        print("- 'help': Show these commands")
        
        while True:
            query_input = input("\nQuestion: ").strip()
            
            if not query_input:
                continue
            
            if query_input.lower() == 'exit':
                break
            elif query_input.lower() == 'help':
                print("\nAvailable commands:")
                print("- 'exit': Quit the program")
                print("- 'switch model': Change the language model")
                print("- 'info': Show ingested documents information")
                print("- 'help': Show these commands")
                print("- 'create podcast': Generate a podcast from ingested content")

            elif query_input.lower() == 'create podcast':
                if not successful_results:
                    print("No documents available for podcast creation!")
                    continue
                
                # Get podcast topic from user
                podcast_query = input("Enter podcast topic or query: ").strip()
                
                try:
                    # Get relevant content from RAG
                    podcast_content = await rag.query(podcast_query)
                    
                    # Generate transcript
                    transcript = await podcast_processor.process_podcast(podcast_content)
                    print("\nGenerated Transcript:\n")
                    print(transcript)
                    
                    # Voice selection
                    print("\nAvailable Voices:")
                    for voice in VOICES:
                        print(f"- {voice}")
                    
                    voice1 = input("Enter name for Host 1 voice: ").strip()
                    voice2 = input("Enter name for Host 2 voice: ").strip()
                    
                    print("Generating podcast audio... hold tight!")

                    # Generate audio
                    audio_url = podcast_processor.generate_audio(
                        transcript=transcript,
                        host1_voice=voice1,
                        host2_voice=voice2
                    )
                    print(f"\nPodcast audio generated successfully! URL: {audio_url}")
                    
                except Exception as e:
                    print(f"Podcast generation failed: {str(e)}")
            elif query_input.lower() == 'info':
                print("\n" + rag.get_ingested_documents_info())
            elif query_input.lower() == 'switch model':
                print("\nAvailable Models:")
                for model in model_manager.list_models():
                    print(f"- {model['name']} ({model['provider']}) - Cost: ${model['cost_per_million_tokens']} per million tokens")
                
                while True:
                    new_model = input("\nEnter new model name: ").strip()
                    try:
                        rag = RAGSystem(new_model)
                        # Reingest documents with new model
                        await rag.ingest_documents(successful_results)
                        print(f"\nSwitched to model: {new_model}")
                        break
                    except ValueError as e:
                        print(f"Error: {e}")
            else:
                try:
                    response = await rag.query(query_input)
                    print(f"\nAnswer: {response}\n")
                except Exception as e:
                    print(f"\nError: Failed to process query - {str(e)}")
                    logging.error("Query processing error", exc_info=True)

    except Exception as e:
        print(f"\nCritical error: {str(e)}")
        logging.error("Critical error in main", exc_info=True)
    finally:
        # Cleanup temp directory
        try:
            for file in os.listdir(processor.tmp_dir):
                os.remove(os.path.join(processor.tmp_dir, file))
            os.rmdir(processor.tmp_dir)
        except Exception as e:
            logging.warning(f"Failed to clean up temporary directory: {str(e)}")

if __name__ == "__main__":
    asyncio.run(main())