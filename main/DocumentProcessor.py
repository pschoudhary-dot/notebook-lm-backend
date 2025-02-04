from datetime import datetime
import asyncio
import platform     
import os
import tempfile
import mimetypes
from typing import Dict
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, CSVLoader, UnstructuredExcelLoader
import aiohttp
from crawl4ai import AsyncWebCrawler
from crawl4ai.async_configs import BrowserConfig, CrawlerRunConfig
from moviepy.editor import VideoFileClip
from urllib.parse import urlparse
import logging
from ProcessingResult import ProcessingResult

class DocumentProcessor:
    def __init__(self):
        self.headers = {"Authorization": f"Bearer {os.getenv('HF_API_KEY')}"}

        if not self.headers["Authorization"]:
            raise ValueError("HF_API_KEY not found in environment variables")
            
        self.browser_config = BrowserConfig(verbose=False)
        self.crawler_run_config = CrawlerRunConfig(
            word_count_threshold=50,
            remove_overlay_elements=True,       
            process_iframes=True
        )
        self.MAX_FILE_SIZE = 25 * 1024 * 1024  # 25MB for audio/video
        self.MAX_AUDIO_LENGTH = 300  # 5 minutes in seconds
        
        # Create tmp directory for processing
        self.tmp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tmp")
        os.makedirs(self.tmp_dir, exist_ok=True)
        
        # Initialize supported mime types
        self.supported_mime_types = {
            'pdf': 'application/pdf',
            'csv': 'text/csv',
            'excel': 'application/vnd.ms-excel',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'doc': 'application/msword',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain',
            'image': 'image/',
            'video': 'video/',
            'audio': 'audio/'
        }

    async def _make_api_request(self, url: str, data: bytes, timeout: int = 300) -> Dict:
        async with aiohttp.ClientSession() as session:
            try:
                async with session.post(url, headers=self.headers, data=data, timeout=timeout) as response:
                    if response.status == 413:
                        raise ValueError("File is too large for API processing")
                    return await response.json()
            except asyncio.TimeoutError:
                raise TimeoutError("API request timed out")

    def _check_file_size(self, file_path: str) -> bool:
        if not os.path.isfile(file_path):
            return True
        
        mime_type = self._get_mime_type(file_path)
        if mime_type and (mime_type.startswith('audio/') or mime_type.startswith('video/')):
            size_limit = self.MAX_FILE_SIZE
        else:
            size_limit = self.MAX_FILE_SIZE * 2
            
        file_size = os.path.getsize(file_path)
        return file_size <= size_limit

    def _get_mime_type(self, file_path: str) -> str:
        mime_type = mimetypes.guess_type(file_path)[0]
        if mime_type is None:
            # Try to determine type from file extension
            ext = os.path.splitext(file_path)[1].lower()
            mime_map = {
                '.pdf': 'application/pdf',
                '.csv': 'text/csv',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel',
                '.doc': 'application/msword',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.txt': 'text/plain',
                '.mp3': 'audio/mpeg',
                '.wav': 'audio/wav',
                '.mp4': 'video/mp4',
                '.avi': 'video/x-msvideo'
            }
            mime_type = mime_map.get(ext)
        return mime_type

    async def process_file(self, file_path: str) -> ProcessingResult:
        try:
            if urlparse(file_path).scheme in ('http', 'https'):
                result = await self._process_url(file_path)
                logging.info(f"Successfully processed URL: {file_path}")
                return ProcessingResult(
                    success=True,
                    content=result,
                    metadata={'type': 'url', 'path': file_path}
                )
            
            if not os.path.exists(file_path):
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"File not found: {file_path}"
                )

            if not self._check_file_size(file_path):
                msg = f"File {file_path} exceeds size limit"
                logging.warning(msg)
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=msg
                )

            mime_type = self._get_mime_type(file_path)
            if not mime_type:
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"Unknown file type: {file_path}"
                )

            logging.info(f"Processing {file_path} ({mime_type})")

            result = None
            metadata = {'type': mime_type, 'path': file_path}

            if mime_type == "application/pdf":
                result = self._process_pdf(file_path)
            elif mime_type in ["text/csv", "application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                result = self._process_csv_excel(file_path)
            elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                result = self._process_doc(file_path)
            elif mime_type and mime_type.startswith('image/'):
                result = await self._process_image(file_path)
            elif mime_type and mime_type.startswith('video/'):
                result = await self._process_video(file_path)
            elif mime_type and mime_type.startswith('audio/'):
                result = await self._process_audio(file_path)
            elif mime_type == "text/plain":
                with open(file_path, 'r', encoding='utf-8') as f:
                    result = f.read()
            else:
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"Unsupported file type: {mime_type}"
                )

            if result:
                logging.info(f"Successfully processed: {file_path}")
                return ProcessingResult(
                    success=True,
                    content=result,
                    metadata=metadata
                )
            else:
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"Failed to process {os.path.basename(file_path)}"
                )

        except Exception as e:
            logging.error(f"Error processing {file_path}: {str(e)}", exc_info=True)
            return ProcessingResult(
                success=False,
                content="",
                error_message=f"Error processing {os.path.basename(file_path)}: {str(e)}"
            )

    def _process_pdf(self, path: str) -> str:
        try:
            loader = PyPDFLoader(path)
            pages = loader.load()
            return "\n".join([p.page_content for p in pages])
        except Exception as e:
            raise RuntimeError(f"Error processing PDF file: {str(e)}")

    def _process_csv_excel(self, path: str) -> str:
        try:
            if path.endswith('.csv'):
                loader = CSVLoader(path)
            else:
                loader = UnstructuredExcelLoader(path)
            docs = loader.load()
            return "\n".join([d.page_content for d in docs])
        except Exception as e:
            raise RuntimeError(f"Error processing CSV/Excel file: {str(e)}")

    def _process_doc(self, path: str) -> str:
        try:
            doc = DocxDocument(path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    full_text.append(para.text)
            # Also process tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            return "\n".join(full_text)
        except Exception as e:
            raise RuntimeError(f"Error processing document file: {str(e)}")

    async def _process_image(self, path: str) -> str:
        API_URL = "https://api-inference.huggingface.co/models/Salesforce/blip-image-captioning-large"
        try:
            with open(path, "rb") as f:
                data = f.read()
            result = await self._make_api_request(API_URL, data)
            return result[0]['generated_text'] if isinstance(result, list) else result.get('error', 'Image processing failed')
        except Exception as e:
            raise RuntimeError(f"Error processing image file: {str(e)}")

    async def _process_video(self, path: str) -> str:
        try:
            video = VideoFileClip(path)
            if video.duration > self.MAX_AUDIO_LENGTH:
                video.close()
                return f"Error: Video duration exceeds {self.MAX_AUDIO_LENGTH} seconds limit"
            
            with tempfile.NamedTemporaryFile(suffix=".wav", dir=self.tmp_dir, delete=False) as tmpfile:
                video.audio.write_audiofile(tmpfile.name, logger=None)
                result = await self._process_audio(tmpfile.name)
                video.close()
                
                try:
                    os.unlink(tmpfile.name)
                except Exception as e:
                    logging.warning(f"Failed to delete temporary file {tmpfile.name}: {str(e)}")
                    
                return f"Video transcription:\n{result}"
        except Exception as e:
            raise RuntimeError(f"Error processing video file: {str(e)}")

    async def _process_audio(self, path: str) -> str:
        API_URL = "https://api-inference.huggingface.co/models/openai/whisper-large-v3-turbo"
        try:
            if os.path.getsize(path) > self.MAX_FILE_SIZE:
                return "Error: Audio file is too large for processing"
                
            with open(path, "rb") as f:
                data = f.read()
            result = await self._make_api_request(API_URL, data)
            transcription = result.get('text', '')
            if not transcription:
                raise ValueError("No transcription received from API")
            return f"Audio transcription:\n{transcription}"
        except Exception as e:
            raise RuntimeError(f"Error processing audio file: {str(e)}")

    async def _process_url(self, path: str) -> str:
        try:
            async with AsyncWebCrawler(config=self.browser_config) as crawler:
                result = await crawler.arun(
                    url=path,
                    config=self.crawler_run_config
                )
                if result.success:
                    return result.markdown
                else:
                    raise RuntimeError("Failed to crawl URL")
        except Exception as e:
            raise RuntimeError(f"Error processing URL: {str(e)}")